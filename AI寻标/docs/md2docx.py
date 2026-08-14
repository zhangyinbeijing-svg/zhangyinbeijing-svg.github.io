# -*- coding: utf-8 -*-
"""将 AI寻标系统PRD-V1.0.md 转换为排版规范的 Word 文档"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'd:\项目\AI寻标\docs\AI寻标系统PRD-V1.0.md'
DST = r'd:\项目\AI寻标\docs\AI寻标系统PRD-V2.0.docx'

BODY_FONT = '宋体'
HEAD_FONT = '微软雅黑'
MONO_FONT = 'Consolas'
HEAD_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x59, 0x59, 0x59)


def set_run(run, size=10.5, bold=False, mono=False, color=None):
    run.font.name = MONO_FONT if mono else BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), MONO_FONT if mono else BODY_FONT)
    if color is not None:
        run.font.color.rgb = color


def add_runs(p, text, size=10.5, color=None):
    """解析 **加粗** 与 `行内代码` 后输出 runs"""
    for m in re.finditer(r'(\*\*.+?\*\*|`.+?`)', text):
        seg = m.group(0)
        if seg.startswith('**'):
            set_run(p.add_run(seg[2:-2]), size=size, bold=True, color=color)
        elif seg.startswith('`'):
            set_run(p.add_run(seg[1:-1]), size=size, mono=True, color=color)
    # 先输出分段:上面只输出了匹配段,需补齐普通段
    return p


def fill_paragraph(p, text, size=10.5, color=None):
    pos = 0
    for m in re.finditer(r'(\*\*.+?\*\*|`.+?`)', text):
        if m.start() > pos:
            set_run(p.add_run(text[pos:m.start()]), size=size, color=color)
        seg = m.group(0)
        if seg.startswith('**'):
            set_run(p.add_run(seg[2:-2]), size=size, bold=True, color=color)
        else:
            set_run(p.add_run(seg[1:-1]), size=size, mono=True, color=color)
        pos = m.end()
    if pos < len(text):
        set_run(p.add_run(text[pos:]), size=size, color=color)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def shade_paragraph(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def init_doc():
    doc = Document()
    # A4 页面
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin, sec.bottom_margin = Cm(2.4), Cm(2.4)
    # 正文样式
    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)
    # 标题样式
    for name, size in [('Heading 1', 17), ('Heading 2', 14),
                       ('Heading 3', 12.5), ('Heading 4', 11.5)]:
        st = doc.styles[name]
        st.font.name = HEAD_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = HEAD_COLOR
        st._element.rPr.rFonts.set(qn('w:eastAsia'), HEAD_FONT)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
    return doc


def flush_table(doc, rows):
    if not rows:
        return
    n_col = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_col)
    tbl.style = 'Table Grid'
    tbl.autofit = True
    for i, row in enumerate(rows):
        for j in range(n_col):
            cell = tbl.cell(i, j)
            txt = row[j] if j < len(row) else ''
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            fill_paragraph(p, txt, size=9.5,
                           color=None)
            if i == 0:
                for r in p.runs:
                    r.font.bold = True
                shade_cell(cell, 'D9E2F3')
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def convert():
    with open(SRC, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = init_doc()
    i, n = 0, len(lines)
    title_done = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1  # 跳过收尾 ```
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            shade_paragraph(p, 'F2F2F2')
            set_run(p.add_run('\n'.join(buf)), size=9.5, mono=True)
            continue

        # 表格
        if stripped.startswith('|'):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if not all(re.match(r'^:?-{2,}:?$', c) for c in cells):
                    rows.append(cells)
                i += 1
            flush_table(doc, rows)
            continue

        # 标题
        m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if m:
            level, text = len(m.group(1)), m.group(2)
            if level == 1 and not title_done:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(18)
                set_run(p.add_run(text), size=20, bold=True, color=HEAD_COLOR)
                title_done = True
            else:
                lvl = min(level, 4)
                p = doc.add_paragraph(style=f'Heading {lvl}')
                fill_paragraph(p, text)
                for r in p.runs:
                    r.font.name = HEAD_FONT
                    r._element.get_or_add_rPr().get_or_add_rFonts().set(
                        qn('w:eastAsia'), HEAD_FONT)
                    r.font.color.rgb = HEAD_COLOR
            i += 1
            continue

        # 水平线
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # 引用
        if stripped.startswith('>'):
            text = re.sub(r'^>\s?', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            shade_paragraph(p, 'EFF3FA')
            fill_paragraph(p, text, color=RGBColor(0x33, 0x33, 0x33))
            i += 1
            continue

        # 无序列表
        m = re.match(r'^-\s+(.*)$', stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            set_run(p.add_run('• '), size=10.5)
            fill_paragraph(p, m.group(1))
            i += 1
            continue

        # 有序列表
        m = re.match(r'^(\d+)[.、]\s*(.*)$', stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            set_run(p.add_run(f'{m.group(1)}. '), size=10.5, bold=True)
            fill_paragraph(p, m.group(2))
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        fill_paragraph(p, stripped)
        i += 1

    doc.core_properties.title = 'AI寻标系统产品需求文档（PRD）'
    doc.core_properties.author = '产品经理'
    doc.save(DST)
    print('saved:', DST)


if __name__ == '__main__':
    convert()
